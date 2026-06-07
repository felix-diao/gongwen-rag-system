import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import List, Dict
import re
from app.config import settings
from app.utils.logger import get_logger

logger = get_logger("text_processor")

class TextProcessor:
    """文本处理工具"""
    
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
    
    def extract_text(self, file_path: str) -> str:
        """从文件提取文本"""
        logger.info(f"开始解析文件: {file_path}")
        if file_path.endswith('.txt') or file_path.endswith('.md'):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif file_path.endswith('.docx'):
            try:
                from docx import Document
                doc = Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                logger.info(f"解析文件成功: {file_path} (docx)")
                return text
            except ImportError:
                raise ValueError("需要安装 python-docx: pip install python-docx")
        
        elif file_path.endswith('.doc'):
            # 旧版 .doc 通过 LibreOffice 转 .docx 再解析
            return self._extract_doc_via_libreoffice(file_path)

        elif file_path.endswith('.pdf'):
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = '\n'.join([page.extract_text() for page in reader.pages])
                    logger.info(f"解析文件成功: {file_path} (pdf)")
                    return text
            except ImportError:
                raise ValueError("需要安装 PyPDF2: pip install PyPDF2")
        
        else:
            logger.warning(f"不支持的文件格式: {file_path}")
            raise ValueError(f"不支持的文件格式: {file_path}")
    
    def _extract_doc_via_libreoffice(self, file_path: str) -> str:
        """通过 LibreOffice headless 将 .doc 转为 .docx 再提取文本。"""
        logger.info(f"尝试 LibreOffice 转换 .doc: {file_path}")
        tmp_dir = None
        try:
            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            if not soffice:
                raise RuntimeError("未找到 LibreOffice，无法解析 .doc 文件，请安装 LibreOffice")
            tmp_dir = tempfile.mkdtemp()
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp_dir, file_path],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
            )
            src_name = Path(file_path).stem
            converted = Path(tmp_dir) / f"{src_name}.docx"
            if not converted.exists():
                raise RuntimeError(f"LibreOffice 转换后未找到输出文件: {converted}")
            from docx import Document
            doc = Document(str(converted))
            text = '\n'.join([para.text for para in doc.paragraphs])
            logger.info(f"LibreOffice 转换 .doc 成功: {file_path}，提取 {len(text)} 字符")
            return text
        except ImportError:
            raise RuntimeError("需要安装 python-docx: pip install python-docx")
        finally:
            if tmp_dir and os.path.isdir(tmp_dir):
                shutil.rmtree(tmp_dir, ignore_errors=True)

    def split_text(self, text: str) -> List[Dict]:
        """智能分块"""
        paragraphs = re.split(r'\n\s*\n', text)
        
        chunks = []
        current_chunk = ""
        current_length = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_length = len(para)
            
            if para_length > self.chunk_size:
                if current_chunk:
                    chunks.append({"text": current_chunk})
                    current_chunk = ""
                    current_length = 0
                
                sub_chunks = self._split_long_paragraph(para)
                chunks.extend(sub_chunks)
                
            elif current_length + para_length > self.chunk_size:
                chunks.append({"text": current_chunk})
                current_chunk = para
                current_length = para_length
                
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
                current_length += para_length
        
        if current_chunk:
            chunks.append({"text": current_chunk})
        
        if self.chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._add_overlap(chunks)
        
        formatted_chunks = []
        for i, chunk in enumerate(chunks):
            formatted_chunks.append({
                "chunk_content": chunk["text"],
                "chunk_index": i
            })
        
        return formatted_chunks
    
    def _split_long_paragraph(self, para: str) -> List[Dict]:
        """分割长段落"""
        sentences = re.split(r'([。！？\.\!\?])', para)
        
        chunks = []
        current_chunk = ""
        current_length = 0
        
        for i in range(0, len(sentences), 2):
            sentence = sentences[i]
            if i + 1 < len(sentences):
                sentence += sentences[i + 1]
            
            sentence = sentence.strip()
            if not sentence:
                continue
            
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size:
                if current_chunk:
                    chunks.append({"text": current_chunk})
                current_chunk = sentence
                current_length = sentence_length
            else:
                current_chunk += sentence
                current_length += sentence_length
        
        if current_chunk:
            chunks.append({"text": current_chunk})
        
        return chunks
    
    def _add_overlap(self, chunks: List[Dict]) -> List[Dict]:
        """添加重叠"""
        overlapped = []
        
        for i, chunk in enumerate(chunks):
            text = chunk["text"]
            
            if i > 0:
                prev_text = chunks[i-1]["text"]
                overlap_text = prev_text[-self.chunk_overlap:]
                text = overlap_text + "\n" + text
            
            overlapped.append({"text": text})
        
        return overlapped
