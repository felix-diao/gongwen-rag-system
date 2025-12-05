import json
import logging
import re
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

from docx import Document
from sqlalchemy.orm import Session

from app.llm_client.generators import get_client
from app.models import database, schemas2
from app.services.meeting_service import file_service
from app.utils.text_processor import TextProcessor

logger = logging.getLogger(__name__)


class MinutesService:
    def __init__(self) -> None:
        self._text_processor = TextProcessor()

    def generate_structured_minutes(
        self,
        db: Session,
        meeting_id: int,
        selected_file_ids: Optional[List[int]] = None,
        audio_segments: Optional[Sequence[Dict[str, Any]]] = None,
    ) -> Optional[schemas2.MeetingInsightsResponse]:
        meeting, combined_text = self._build_meeting_context(
            db,
            meeting_id,
            selected_file_ids,
            audio_segments,
        )
        if not meeting:
            return None

        payload = self._invoke_structured_llm(combined_text, meeting)
        summary_text = (payload or {}).get("summary") or self._fallback_summary(meeting)
        action_items_raw = (payload or {}).get("action_items")
        decision_items_raw = (payload or {}).get("decision_items")

        normalized_actions = self._normalize_action_items(action_items_raw, meeting)
        normalized_decisions = self._normalize_decision_items(decision_items_raw, meeting)

        self._reset_structured_minutes(db, meeting_id)

        summary_record = database.MeetingSummary(
            meeting_id=meeting_id,
            summary_text=summary_text,
        )
        db.add(summary_record)

        action_records: List[database.MeetingActionItem] = []
        for item in normalized_actions:
            action = database.MeetingActionItem(
                meeting_id=meeting_id,
                description=item.get("description"),
                owner=item.get("owner"),
                due_date=item.get("due_date"),
                status=item.get("status") or "pending",
            )
            db.add(action)
            action_records.append(action)

        decision_records: List[database.MeetingDecisionItem] = []
        for item in normalized_decisions:
            decision = database.MeetingDecisionItem(
                meeting_id=meeting_id,
                description=item.get("description"),
            )
            db.add(decision)
            decision_records.append(decision)

        db.commit()
        db.refresh(summary_record)
        for record in action_records + decision_records:
            db.refresh(record)

        return schemas2.MeetingInsightsResponse(
            summary=schemas2.MeetingSummaryInDB.from_orm(summary_record),
            action_items=[schemas2.MeetingActionItemInDB.from_orm(item) for item in action_records],
            decision_items=[schemas2.MeetingDecisionItemInDB.from_orm(item) for item in decision_records],
        )

    def get_meeting_insights(self, db: Session, meeting_id: int) -> Optional[schemas2.MeetingInsightsResponse]:
        summary = self._get_summary_record(db, meeting_id)
        actions = self.list_action_items(db, meeting_id)
        decisions = self.list_decision_items(db, meeting_id)

        if not summary and not actions and not decisions:
            return None

        return schemas2.MeetingInsightsResponse(
            summary=schemas2.MeetingSummaryInDB.from_orm(summary) if summary else None,
            action_items=[schemas2.MeetingActionItemInDB.from_orm(item) for item in actions],
            decision_items=[schemas2.MeetingDecisionItemInDB.from_orm(item) for item in decisions],
        )

    def get_summary(self, db: Session, meeting_id: int):
        summary = self._get_summary_record(db, meeting_id)
        return summary

    def update_summary(self, db: Session, meeting_id: int, summary_update: schemas2.MeetingSummaryUpdate):
        summary = self._get_summary_record(db, meeting_id)
        new_text = summary_update.summary_text

        if summary is None:
            if not new_text:
                meeting = self._get_meeting(db, meeting_id)
                new_text = self._fallback_summary(meeting) if meeting else "会议摘要待补充"
            summary = database.MeetingSummary(meeting_id=meeting_id, summary_text=new_text)
            db.add(summary)
        else:
            if new_text is not None:
                summary.summary_text = new_text

        db.commit()
        db.refresh(summary)
        return summary

    def list_action_items(self, db: Session, meeting_id: int) -> List[database.MeetingActionItem]:
        return (
            db.query(database.MeetingActionItem)
            .filter(database.MeetingActionItem.meeting_id == meeting_id)
            .order_by(database.MeetingActionItem.id.asc())
            .all()
        )

    def get_action_item(self, db: Session, meeting_id: int, item_id: int) -> Optional[database.MeetingActionItem]:
        return (
            db.query(database.MeetingActionItem)
            .filter(
                database.MeetingActionItem.meeting_id == meeting_id,
                database.MeetingActionItem.id == item_id,
            )
            .first()
        )

    def create_action_item(
        self,
        db: Session,
        meeting_id: int,
        action_data: schemas2.MeetingActionItemCreate,
    ) -> database.MeetingActionItem:
        payload = action_data.dict(exclude_unset=True)
        action = database.MeetingActionItem(
            meeting_id=meeting_id,
            description=payload.get("description"),
            owner=payload.get("owner"),
            due_date=payload.get("due_date"),
            status=payload.get("status") or "pending",
        )
        db.add(action)
        db.commit()
        db.refresh(action)
        return action

    def update_action_item(
        self,
        db: Session,
        meeting_id: int,
        item_id: int,
        action_update: schemas2.MeetingActionItemUpdate,
    ) -> Optional[database.MeetingActionItem]:
        action = self.get_action_item(db, meeting_id, item_id)
        if not action:
            return None
        for key, value in action_update.dict(exclude_unset=True).items():
            setattr(action, key, value)
        db.commit()
        db.refresh(action)
        return action

    def delete_action_item(self, db: Session, meeting_id: int, item_id: int) -> bool:
        action = self.get_action_item(db, meeting_id, item_id)
        if not action:
            return False
        db.delete(action)
        db.commit()
        return True

    def list_decision_items(self, db: Session, meeting_id: int) -> List[database.MeetingDecisionItem]:
        return (
            db.query(database.MeetingDecisionItem)
            .filter(database.MeetingDecisionItem.meeting_id == meeting_id)
            .order_by(database.MeetingDecisionItem.id.asc())
            .all()
        )

    def get_decision_item(self, db: Session, meeting_id: int, item_id: int) -> Optional[database.MeetingDecisionItem]:
        return (
            db.query(database.MeetingDecisionItem)
            .filter(
                database.MeetingDecisionItem.meeting_id == meeting_id,
                database.MeetingDecisionItem.id == item_id,
            )
            .first()
        )

    def create_decision_item(
        self,
        db: Session,
        meeting_id: int,
        decision_data: schemas2.MeetingDecisionItemCreate,
    ) -> database.MeetingDecisionItem:
        payload = decision_data.dict(exclude_unset=True)
        decision = database.MeetingDecisionItem(
            meeting_id=meeting_id,
            description=payload.get("description"),
        )
        db.add(decision)
        db.commit()
        db.refresh(decision)
        return decision

    def update_decision_item(
        self,
        db: Session,
        meeting_id: int,
        item_id: int,
        decision_update: schemas2.MeetingDecisionItemUpdate,
    ) -> Optional[database.MeetingDecisionItem]:
        decision = self.get_decision_item(db, meeting_id, item_id)
        if not decision:
            return None
        updates = decision_update.dict(exclude_unset=True)
        description = updates.get("description")
        if description is not None:
            decision.description = description
        db.commit()
        db.refresh(decision)
        return decision

    def delete_decision_item(self, db: Session, meeting_id: int, item_id: int) -> bool:
        decision = self.get_decision_item(db, meeting_id, item_id)
        if not decision:
            return False
        db.delete(decision)
        db.commit()
        return True

    def export_structured_docx(self, db: Session, meeting_id: int) -> Optional[Path]:
        meeting = self._get_meeting(db, meeting_id)
        if not meeting:
            return None

        summary = self._get_summary_record(db, meeting_id)
        actions = self.list_action_items(db, meeting_id)
        decisions = self.list_decision_items(db, meeting_id)

        summary_text = summary.summary_text if summary else self._fallback_summary(meeting)

        export_dir = self._get_export_dir(meeting_id)
        export_dir.mkdir(parents=True, exist_ok=True)
        filename = f"structured_minutes_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.docx"
        file_path = export_dir / filename

        self._write_structured_docx(file_path, meeting, summary_text, actions, decisions)

        self._save_export_file(
            db,
            meeting_id,
            file_path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return file_path

    def _reset_structured_minutes(self, db: Session, meeting_id: int) -> None:
        db.query(database.MeetingSummary).filter(database.MeetingSummary.meeting_id == meeting_id).delete(synchronize_session=False)
        db.query(database.MeetingActionItem).filter(database.MeetingActionItem.meeting_id == meeting_id).delete(synchronize_session=False)
        db.query(database.MeetingDecisionItem).filter(database.MeetingDecisionItem.meeting_id == meeting_id).delete(synchronize_session=False)
        db.flush()

    def _get_summary_record(self, db: Session, meeting_id: int) -> Optional[database.MeetingSummary]:
        return (
            db.query(database.MeetingSummary)
            .filter(database.MeetingSummary.meeting_id == meeting_id)
            .first()
        )

    def _get_meeting(self, db: Session, meeting_id: int) -> Optional[database.Meeting]:
        return db.query(database.Meeting).filter(database.Meeting.id == meeting_id).first()

    def _get_export_dir(self, meeting_id: int) -> Path:
        repo_root = Path(__file__).resolve().parents[2]
        return repo_root / "meeting_files" / str(meeting_id) / "exports"

    def _build_meeting_context(
        self,
        db: Session,
        meeting_id: int,
        selected_file_ids: Optional[List[int]],
        audio_segments: Optional[Sequence[Dict[str, Any]]],
    ):
        meeting = self._get_meeting(db, meeting_id)
        if not meeting:
            return None, ""

        meeting_info = [f"会议标题: {meeting.title}"]
        if meeting.date:
            meeting_info.append(f"会议时间: {meeting.date}")
        if meeting.location:
            meeting_info.append(f"会议地点: {meeting.location}")
        if meeting.host:
            meeting_info.append(f"主持人: {meeting.host}")
        if meeting.participants:
            meeting_info.append(f"参会人员: {meeting.participants}")

        parts = ["\n".join(meeting_info)]
        if getattr(meeting, "content_text", None):
            parts.append(f"【会议记录】\n{meeting.content_text}")

        if selected_file_ids:
            logger.info("生成纪要时使用所选文件ids: %s", selected_file_ids)
            files_query = db.query(database.MeetingFile).filter(
                database.MeetingFile.meeting_id == meeting_id,
                database.MeetingFile.id.in_(selected_file_ids),
            )
        else:
            files_query = db.query(database.MeetingFile).filter(database.MeetingFile.meeting_id == meeting_id)

        files = files_query.all()
        for idx, f in enumerate(files, start=1):
            fp = getattr(f, "file_path", None)
            if not fp:
                continue
            try:
                text = self._text_processor.extract_text(fp)
                parts.append(f"【文件{idx}:{f.filename}】\n{text}")
                logger.info("提取文件文本成功: %s", fp)
            except Exception as exc:
                logger.warning("提取文件文本失败: %s，错误: %s", fp, exc)
                parts.append(f"【文件{idx}:{f.filename}】\n(无法提取文本)")

        if audio_segments:
            for idx, segment in enumerate(audio_segments, start=1):
                if not isinstance(segment, dict):
                    continue
                name = segment.get("name") or f"音频{idx}"
                text = segment.get("text")
                if not text:
                    continue
                parts.append(f"【音频{idx}:{name}】\n{text}")

        combined_text = "\n\n".join(parts)
        logger.debug("会议数据汇总完成，会议ID: %s，文本长度: %s", meeting_id, len(combined_text))
        return meeting, combined_text

    def _invoke_structured_llm(self, combined_text: str, meeting: database.Meeting) -> Optional[dict]:
        instruction = (
            "你是一名会议助理，需要根据提供的原始资料生成 JSON，字段包括 summary, action_items, decision_items。"
            "summary 为一句话概述；action_items 是数组，每一项包含 description、owner、due_date(YYYY-MM-DD，可为空)、status；"
            "decision_items 为数组，每一项包含 description。"
        )
        user_msg = (
            f"请输出严格的 JSON，不要包含额外说明。\n会议标题: {meeting.title}\n"
            f"原始资料如下:\n{combined_text}"
        )
        try:
            cli = get_client()
            raw = cli.chat(
                [
                    {"role": "system", "content": instruction},
                    {"role": "user", "content": user_msg},
                ],
                max_tokens=1500,
            )
            return self._parse_structured_payload(raw)
        except Exception as exc:
            logger.warning("生成结构化纪要失败，错误: %s", exc)
            return None

    def _parse_structured_payload(self, raw: str) -> Optional[dict]:
        if not raw:
            return None
        cleaned = raw.strip()
        match = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if match:
            cleaned = match.group(0)
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            logger.warning("解析结构化纪要 JSON 失败，原始内容: %s", raw)
            return None

    def _fallback_summary(self, meeting: database.Meeting) -> str:
        base = meeting.title or "会议"
        return f"本次《{base}》会议聚焦核心议题，团队明确了重点任务与下一阶段推进计划。"

    def _normalize_action_items(self, raw_items: Any, meeting: database.Meeting) -> List[Dict[str, Any]]:
        normalized: List[Dict[str, Any]] = []
        if isinstance(raw_items, list):
            for item in raw_items:
                if not isinstance(item, dict):
                    continue
                description = (item.get("description") or item.get("action") or "").strip()
                if not description:
                    continue
                normalized.append(
                    {
                        "description": description,
                        "owner": (item.get("owner") or item.get("responsible") or "").strip() or None,
                        "due_date": self._parse_due_date(item.get("due_date")),
                        "status": (item.get("status") or "pending").strip() or "pending",
                    }
                )

        if not normalized:
            normalized.append(
                {
                    "description": f"根据《{meeting.title}》会议讨论完善详细行动项",
                    "owner": meeting.host or "待指派",
                    "due_date": None,
                    "status": "pending",
                }
            )
        return normalized

    def _normalize_decision_items(self, raw_items: Any, meeting: database.Meeting) -> List[Dict[str, Any]]:
        normalized: List[Dict[str, Any]] = []
        if isinstance(raw_items, list):
            for item in raw_items:
                if isinstance(item, dict):
                    description = (item.get("description") or item.get("decision") or "").strip()
                else:
                    description = str(item).strip()
                if description:
                    normalized.append({"description": description})

        if not normalized:
            normalized.append({"description": f"会议确认推进《{meeting.title}》后续实施计划"})
        return normalized

    def _parse_due_date(self, value: Any) -> Optional[date]:
        if not value:
            return None
        if isinstance(value, date) and not isinstance(value, datetime):
            return value
        if isinstance(value, datetime):
            return value.date()
        if isinstance(value, str):
            candidate = value.strip().split()[0]
            for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y%m%d"):
                try:
                    return datetime.strptime(candidate, fmt).date()
                except ValueError:
                    continue
        return None

    def _write_structured_docx(
        self,
        output_path: Path,
        meeting: database.Meeting,
        summary_text: str,
        actions: List[database.MeetingActionItem],
        decisions: List[database.MeetingDecisionItem],
    ) -> None:
        document = Document()
        document.add_heading(meeting.title or "会议纪要", level=1)

        info_lines = []
        if meeting.date:
            info_lines.append(f"会议时间: {meeting.date.strftime('%Y-%m-%d %H:%M')}")
        if meeting.location:
            info_lines.append(f"会议地点: {meeting.location}")
        if meeting.host:
            info_lines.append(f"主持人: {meeting.host}")
        if meeting.participants:
            info_lines.append(f"参会人员: {meeting.participants}")
        if info_lines:
            for line in info_lines:
                document.add_paragraph(line)
        document.add_paragraph("")

        document.add_heading("会议摘要", level=2)
        document.add_paragraph(summary_text or "暂无摘要")

        document.add_heading("行动项", level=2)
        if actions:
            for idx, action in enumerate(actions, start=1):
                due = action.due_date.strftime('%Y-%m-%d') if action.due_date else "-"
                owner = action.owner or "待指派"
                status = action.status or "pending"
                document.add_paragraph(
                    f"{idx}. {action.description} (负责人: {owner}，截止: {due}，状态: {status})"
                )
        else:
            document.add_paragraph("暂无行动项")

        document.add_heading("决策事项", level=2)
        if decisions:
            for idx, decision in enumerate(decisions, start=1):
                document.add_paragraph(f"{idx}. {decision.description}")
        else:
            document.add_paragraph("暂无决策事项")

        document.save(output_path)

    def _save_export_file(self, db: Session, meeting_id: int, file_path: Path, mime: str):
        payload = schemas2.MeetingFileCreate(
            meeting_id=meeting_id,
            filename=file_path.name,
            file_type=mime,
            file_path=str(file_path),
        )
        return file_service.create_file(db, payload)



# 创建服务实例
minutes_service = MinutesService()
