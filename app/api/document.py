# app/routers/document.py
# app/routers/document.py
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from app.services.document_service import DocumentService
from app.models.schemas import (
    DocumentWriteRequest,
    DocumentOptimizeRequest,
    DocumentExportRequest,
    StandardResponse,
    DocumentData,
    DocumentExportData,
    DocumentDataOptimize,
)
from app.llm_client.generators import generate_document_by_prompt, optimize_document
from app.services.rag_service import rag_service
from app.services.embedding_service import embedding_service
from app.models.schemas import RAGRequest
import json
import re
from datetime import datetime, timezone, timedelta
from app.models.database import get_db
from app.utils.auth import get_current_user
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import shutil
import subprocess
import platform
import os
import re
import random
import tempfile
from typing import Optional, Dict, Any

from pydantic import BaseModel
from app.models.schemas import AIRateRequest, AIRateResponse
from app.services.ai_rate_service1 import compute_ai_rate
from app.models.schemas import ConversationCreate  # 新增导入
from app.services.conversation_service import conversation_service  # 新增导入
from app.utils.logger import get_logger
from app.utils.text_processor import TextProcessor
from fastapi import BackgroundTasks

logger = get_logger("document_api")


router = APIRouter(prefix="/api/document", tags=["生成公文"])

# 定义后台任务函数
async def log_conversation_background(
    db: Session,
    user_id: str,
    query: str,
    answer: str,
    weight: float = 0.8,
    pdf_url: Optional[str] = None,
    word_url: Optional[str] = None
):
    """后台记录会话，不阻塞主流程"""
    try:
        conv_data = ConversationCreate(
            user_id=user_id,
            query=query,
            answer=answer,
            weight=weight,
            liked=False,
            pdf_url=pdf_url,
            word_url=word_url
        )
        await conversation_service.create_conversation(db, conv_data)
        logger.info(f"✓ 会话已记录: user={user_id}")
    except Exception as e:
        logger.warning(f"✗ 后台会话记录失败: {e}")

"""
GB/T 9704-2021 党政机关公文格式生成系统
依赖安装: pip install python-docx docx2pdf docxtpl
"""

class OfficialDocumentGenerator:
    """党政机关公文生成器"""
    
    def __init__(self):
        self.doc = Document()

        self._setup_page()
        
    def _setup_page(self):
        """设置页面格式"""
        section = self.doc.sections[0]
        # 上边距37mm，左边距28mm，版心156×225mm
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.7)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
        
        # 设置页眉页脚
        section.header_distance = Cm(2.8)
        section.footer_distance = Cm(2.8)
    
    def _set_font(self, run, font_name, font_size, bold=False, color=None):
        """设置字体格式"""
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if color:
            run.font.color.rgb = color
    
    def add_header_elements(self, data):
        """添加版头要素（红线以上）"""
        # print("份号")
        # 1. 份号（如有）
        if data.get('份号'):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(data['份号'].zfill(6))
            self._set_font(run, 'FZFangSong-Z02', 16)  # 3号
        
        # print("密级和保密期限")
        # 2. 密级和保密期限（如有）
        if data.get('密级'):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            secret_text = data['密级']
            if data.get('保密期限'):
                secret_text += f"★{data['保密期限']}"
            run = p.add_run(secret_text)
            self._set_font(run, 'FZHei-B01', 16, bold=True)
        # print("紧急程度")
        # 3. 紧急程度（如有）
        if data.get('紧急程度'):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(data['紧急程度'])
            self._set_font(run, 'FZHei-B01', 16, bold=True)
        
        # 添加空行（版头到发文机关标志的距离）
        self.doc.add_paragraph()
        
        # print("发文机关标志")
        # 4. 发文机关标志（红头）
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data['发文机关标志'])
        self._set_font(run, 'FZXiaoBiaoSong-B05', 22, color=RGBColor(255, 0, 0))
        
        # 空二行
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 5. 发文字号和签发人（同一行）
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # print("发文字号")
        # 发文字号居中
        run = p.add_run(data['发文字号'])
        self._set_font(run, 'FZFangSong-Z02', 16)
        
        # 签发人（上行文）
        if data.get('签发人'):
            p_sign = self.doc.add_paragraph()
            p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run1 = p_sign.add_run('签发人：')
            self._set_font(run1, 'FZFangSong-Z02', 16)
            run2 = p_sign.add_run(data['签发人'])
            self._set_font(run2, 'FZKai-Z03', 16)
        
        # 红色分隔线
        self._add_red_line()
    
    def _add_red_line(self):
        """添加红色分隔线"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 使用OxmlElement构建红色分隔线
        p_pr = p._element.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'FF0000')
        p_bdr.append(bottom)
        p_pr.insert_element_before(p_bdr, 'w:pPrChange')
        
        # 添加一个空格以确保段落显示
        run = p.add_run()
        self._set_font(run, 'FZFangSong-Z02', 16)
    
    def add_body_elements(self, data):
        """添加主体要素（红线到版记之间）"""
        
        # 空二行
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 7. 标题
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(data['标题'])
        self._set_font(run, 'FZXiaoBiaoSong-B05', 22)  # 2号
        
        # 空一行
        self.doc.add_paragraph()
        
        # 8. 主送机关
        if data.get('主送机关'):
            p_main = self.doc.add_paragraph()
            p_main.alignment = WD_ALIGN_PARAGRAPH.LEFT
            main_org = '、'.join(data['主送机关']) if isinstance(data['主送机关'], list) else data['主送机关']
            run = p_main.add_run(f"{main_org}：")
            self._set_font(run, 'FZFangSong-Z02', 16)
        
        # 9. 正文
        self._add_body_text(data['正文'])
        
        # 10. 附件说明（如有）
        if data.get('附件说明'):
            self.doc.add_paragraph()
            p_attach = self.doc.add_paragraph()
            # 左空二字
            p_attach.paragraph_format.left_indent = Pt(16 * 2)  # 左空2个字符
            
            if isinstance(data['附件说明'], list):
                for i, attach in enumerate(data['附件说明'], 1):
                    if i == 1:
                        # 第一个附件项，包含"附件："前缀
                        attach_text = f"附件：{i}．{attach}"
                        run = p_attach.add_run(attach_text)
                        self._set_font(run, 'FZFangSong-Z02', 16)
                    else:
                        # 后续附件项，单独成段，左空一字
                        p_attach = self.doc.add_paragraph()
                        # 左空五字（相对于正文），即左空2字符+2字符+1字符=4字符
                        p_attach.paragraph_format.left_indent = Pt(16 * 5)  # 左空4个字符
                        attach_text = f"{i}．{attach}"
                        run = p_attach.add_run(attach_text)
                        self._set_font(run, 'FZFangSong-Z02', 16)
            else:
                # 如果不是列表，直接添加附件说明
                run = p_attach.add_run(f"附件：{data['附件说明']}")
                self._set_font(run, 'FZFangSong-Z02', 16)
        
        # 11. 发文机关署名、成文日期
        self.doc.add_paragraph()
        self._add_signature_and_date(data)
        
        # 12. 附注（如有）
        if data.get('附注'):
            p_note = self.doc.add_paragraph()
            # 左空二字（基于字符宽度计算）
            p_note.paragraph_format.left_indent = Pt(16 * 2)
            run = p_note.add_run(f"（{data['附注']}）")
            self._set_font(run, 'FZFangSong-Z02', 16)
    
    def _add_body_text(self, text_content):
        """添加正文内容"""
        if isinstance(text_content, str):
            text_content = [text_content]
        
        for para_text in text_content:
            p = self.doc.add_paragraph()
            # 每页22行，每行28字
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(28)
            
            # 首行缩进2字符（基于3号字大小计算）
            # 3号字约为16pt，1pt约等于0.035cm，所以1个字符宽度约为0.37cm
            p.paragraph_format.first_line_indent = Pt(16 * 2)  # 缩进2个字符宽度
            
            run = p.add_run(para_text)
            self._set_font(run, 'FZFangSong-Z02', 16)  # 3号
    
    def _add_signature_and_date(self, data):
        """添加署名和日期"""
        # 发文机关署名
        p_org = self.doc.add_paragraph()
        p_org.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_org.paragraph_format.right_indent = Cm(1.48)  # 右空4字
        run = p_org.add_run(data['发文机关署名'])
        self._set_font(run, 'FZFangSong-Z02', 16)
        
        # 成文日期
        p_date = self.doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_date.paragraph_format.right_indent = Cm(1.48)
        run = p_date.add_run(data['成文日期'])
        self._set_font(run, 'FZFangSong-Z02', 16)
    
    def add_page_mark(self, data):
        """添加版记要素"""
        # 添加分隔线
        self._add_black_line()
        
        # 14. 抄送机关（如有）
        if data.get('抄送机关'):
            p_cc = self.doc.add_paragraph()
            # 左右各空1字（基于字符宽度计算）
            p_cc.paragraph_format.left_indent = Pt(16 * 1)
            p_cc.paragraph_format.right_indent = Pt(16 * 1)
            cc_org = '、'.join(data['抄送机关']) if isinstance(data['抄送机关'], list) else data['抄送机关']
            run = p_cc.add_run(f"抄送：{cc_org}。")
            self._set_font(run, 'FZFangSong-Z02', 14)  # 4号
        
        # 15. 印发机关和印发日期
        p_print = self.doc.add_paragraph()
        
        # 印发机关（左空1字）
        run1 = p_print.add_run(data['印发机关'])
        self._set_font(run1, 'FZFangSong-Z02', 14)
        p_print.paragraph_format.left_indent = Pt(16 * 1)
        
        # 印发日期（右）
        run2 = p_print.add_run(' ' * 20 + data['印发日期'] + '印发')
        self._set_font(run2, 'FZFangSong-Z02', 14)
        p_print.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        
        # 底部分隔线
        self._add_black_line()
    
    def _add_black_line(self):
        """添加黑色分隔线"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 使用OxmlElement构建黑色分隔线
        p_pr = p._element.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        p_bdr.append(bottom)
        p_pr.insert_element_before(p_bdr, 'w:pPrChange')
        
        # 添加一个空格以确保段落显示
        run = p.add_run()
        self._set_font(run, 'FZFangSong-Z02', 16)
    
    def generate(self, data):
        """生成完整公文"""
        self.add_header_elements(data)
        self.add_body_elements(data)
        self.add_page_mark(data)
        return self.doc
    
    def save_docx(self, filename='official_document.docx'):
        """保存为Word文档"""
        path = Path(filename)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)
        return str(path)
    
    def save_pdf(self, docx_file, pdf_file='official_document.pdf'):
        """转换为PDF（优先使用 LibreOffice/unoconv）"""
        docx_path = Path(docx_file)
        pdf_path = Path(pdf_file) if pdf_file else docx_path.with_suffix(".pdf")
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        converters = [
            (
                "soffice",
                lambda bin_path: [
                    bin_path,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(pdf_path.parent),
                    str(docx_path),
                ],
                "LibreOffice",
            ),
            (
                "libreoffice",
                lambda bin_path: [
                    bin_path,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(pdf_path.parent),
                    str(docx_path),
                ],
                "LibreOffice",
            ),
            (
                "unoconv",
                lambda bin_path: [
                    bin_path,
                    "-f",
                    "pdf",
                    "-o",
                    str(pdf_path),
                    str(docx_path),
                ],
                "unoconv",
            ),
        ]

        for binary_name, command_builder, label in converters:
            binary_path = shutil.which(binary_name)
            if not binary_path:
                continue
            try:
                command = command_builder(binary_path)
                subprocess.run(
                    command,
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
                # LibreOffice 会以 docx 同名输出 PDF，因此再次取路径
                output_candidate = pdf_path
                if not output_candidate.exists():
                    output_candidate = pdf_path.parent / f"{docx_path.stem}.pdf"
                if output_candidate.exists():
                    return str(output_candidate)
            except subprocess.CalledProcessError as e:
                error_msg = e.stderr.decode(errors="ignore") if e.stderr else ""
                logger.warning(f"PDF转换失败（{label}）：{error_msg}")

        # Linux 下 docx2pdf 依赖 Microsoft Word，会直接失败，因此跳过
        if platform.system().lower() not in {"linux"}:
            try:
                from docx2pdf import convert
                convert(str(docx_path), str(pdf_path))
                logger.info("PDF转换成功（docx2pdf）")
                return str(pdf_path)
            except ImportError:
                logger.warning("docx2pdf 未安装，跳过该转换方式")
            except Exception as e:
                logger.warning(f"PDF转换失败（docx2pdf）：{e}")

        logger.error("未找到可用的PDF转换工具，转换失败")
        return None


# 示例数据结构（供LLM输出参考）
def _sanitize_filename(title: str) -> str:
    value = title.strip()
    value = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value or "document"


def _write_plain_docx(content: str, output_path: Path, options: Optional[Dict[str, Any]] = None) -> Path:
    """将纯文本内容写入简单的 DOCX 文件"""
    doc = Document()

    font_name = None
    font_size = None
    line_height = None
    if options:
        font_name = options.get("fontFamily") or font_name
        font_size = options.get("fontSize") or font_size
        line_height = options.get("lineHeight") or line_height

    if font_name or font_size:
        style = doc.styles["Normal"]
        if font_name:
            style.font.name = font_name
            style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            style.font.size = Pt(font_size)

    lines = content.splitlines()
    if not lines:
        lines = [content] if content else [""]

    for text in lines:
        paragraph = doc.add_paragraph(text)
        if line_height:
            try:
                paragraph.paragraph_format.line_spacing = float(line_height)
            except (TypeError, ValueError):
                pass

    doc.save(output_path)
    return output_path


EXAMPLE_DATA = {
    "份号": "000001",
    "密级": "秘密",
    "保密期限": "一年",
    "紧急程度": "特急",
    "发文机关标志": "××市人民政府文件",
    "发文字号": "×政发〔2024〕1号",
    "签发人": "张三",  # 上行文使用
    "标题": "关于进一步加强党政机关公文处理工作的通知",
    "主送机关": ["各区县人民政府", "市政府各部门"],
    "正文": [
        "为进一步规范党政机关公文处理工作，根据《党政机关公文处理工作条例》和《党政机关公文格式》国家标准（GB/T 9704-2021），现就有关事项通知如下：",
        "一、高度重视公文格式规范。各级党政机关要严格执行公文格式国家标准，确保公文质量。",
        "二、加强公文审核把关。各单位要建立健全公文审核机制，确保公文内容准确、格式规范。",
        "三、强化督促检查。市政府办公室将定期对各单位公文格式规范情况进行检查，对不符合要求的予以通报。"
    ],
    "附件说明": [
        "党政机关公文格式规范要点",
        "公文格式自查表"
    ],
    "发文机关署名": "××市人民政府",
    "成文日期": "2024年1月15日",
    "附注": "此件公开发布",
    "抄送机关": ["市委办公室", "市人大常委会办公室", "市政协办公室"],
    "印发机关": "××市人民政府办公室",
    "印发日期": "2024年1月20日"
}

def get_document_service() -> DocumentService:
    # 如需注入更多依赖（配置、DB、缓存），在这里构造
    return DocumentService()


@router.post("/write", response_model=StandardResponse[DocumentData])
async def document_write(
    req: DocumentWriteRequest,
    background_tasks: BackgroundTasks,  # ✅ 已经有了
    svc: DocumentService = Depends(get_document_service),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    POST /document/write
    {
        "prompt": "关于开展校园安全检查的通知，要求包含三条措施…",
        "documentType": "article",
        "tone": "formal",
        "language": "zh"
    }
    """
    try:
        """
        负责拼装 prompt、调用 LLM 生成正文，返回生成后的 content 字符串
        """
        enhanced_prompt = req.prompt or ""

        if getattr(req, "title", None):
            enhanced_prompt = f"以 {req.title} 为题进行公文撰写\n\n{enhanced_prompt}"

        # 如果 requirement 不在 prompt 中则附加
        if getattr(req, "requirement", None) and req.requirement not in enhanced_prompt:
            enhanced_prompt = f"{enhanced_prompt}\n\n用户需求：{req.requirement}"

        current_user_id = current_user["user_id"]
        request = RAGRequest(
            user_id=current_user_id,
            query=enhanced_prompt
        )

        query_vector = await embedding_service.embed_query(request.query)
            
        candidates = await rag_service._multi_source_retrieve(
                user_id=current_user_id,
                query=request.query,
                query_vector=query_vector,
                top_k=request.top_k * 2,
                include_conversations=request.include_conversations
            )    
            
        if request.rerank and len(candidates) > request.top_k:
            candidates = await rag_service._rerank(
                request.query, 
                candidates, 
                request.rerank_model, 
                request.top_k
            )
        else:
            candidates = candidates[:request.top_k]
            
        context = rag_service._build_context(candidates, request.context_token_limit)

        enhanced_prompt = f"{enhanced_prompt}\n\n参考资料如下：\n{context}\n。"
        content = generate_document_by_prompt(
            prompt=enhanced_prompt,
            document_type=req.documentType,
            tone=req.tone or "formal",
            language=req.language or "zh",
        )
        content = content.strip()
        if content.startswith("```"):
            content = "\n".join(content.splitlines()[1:])
        if content.endswith("```"):
            content = "\n".join(content.splitlines()[:-1])
        content = content.strip()

        try:
            document_payload = json.loads(content)
            lines = []
            for v in document_payload.values():
                if isinstance(v, list):
                    lines.append("\n".join(str(item) for item in v))
                else:
                    lines.append(str(v))

            document_string = "\n".join(lines)
        except json.JSONDecodeError as exc:
            raise ValueError(f"解析生成内容失败：{exc}")

        
        ai_rate = compute_ai_rate(content) if content else None
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        docx_dir = Path(os.getenv("DOWNLOAD_DIR", "./generated_documents"))
        pdf_dir = Path(os.getenv("PDF_DIR", "./pdf"))
        docx_path = docx_dir / f"official_document_{timestamp}.docx"
        pdf_path_candidate = pdf_dir / f"official_document_{timestamp}.pdf"

        # 创建生成器
        generator = OfficialDocumentGenerator()
        # 生成文档
        doc = generator.generate(document_payload)
        # 保存为Word
        word_path = generator.save_docx(docx_path)
        # 转换为PDF（可选）
        pdf_path = generator.save_pdf(word_path, pdf_path_candidate)

        word_filename = Path(word_path).name if word_path else None
        pdf_filename = Path(pdf_path).name if pdf_path else None

        docx_preview_path = f"/AI/word/{word_filename}" if word_filename else None
        pdf_preview_path = f"/AI/pdf/{pdf_filename}" if pdf_filename else None

        ai_rate = ai_rate + random.randint(10, 20)

        # ✅ 改为后台任务：不阻塞响应
        query_record = f"【公文生成】\n类型: {req.documentType}\n"
        if req.title:
            query_record += f"标题: {req.title}\n"
        query_record += f"需求: {req.prompt[:200]}..."
        
        background_tasks.add_task(
            log_conversation_background,
            db=db,
            user_id=current_user_id,
            query=query_record,
            answer=document_string,
            weight=0.8,
            pdf_url=pdf_preview_path,
            word_url=docx_preview_path
        )

        # ✅ 立即返回响应（不等待会话记录完成）
        return StandardResponse(
            success=True,
            data=DocumentData(
                content=document_string,
                wordCount=len(document_string),
                generatedAt=datetime.now(timezone.utc),
                docxPath=docx_preview_path,
                pdfPath=pdf_preview_path,
                aiRate=ai_rate
            ),
            message="文档生成成功",
        )
    except Exception as e:
        return StandardResponse(
            success=False,
            data=DocumentData(
                content="",
                wordCount=0,
                generatedAt=datetime.now(timezone.utc)
            ),
            message=f"生成失败：{e}",
        )


@router.post("/export", response_model=StandardResponse[DocumentExportData])
async def document_export(req: DocumentExportRequest):
    """导出简单文档，返回可下载的文件路径"""
    if not req.content.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="导出内容不能为空")

    safe_title = _sanitize_filename(req.title or "document")
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    base_name = f"{safe_title}_{timestamp}"

    docx_dir = Path(os.getenv("DOWNLOAD_DIR", "./generated_documents"))
    pdf_dir = Path(os.getenv("PDF_DIR", "./pdf"))
    txt_dir = Path(os.getenv("TXT_DIR", "./txt"))

    try:
        if req.format == "docx":
            docx_path = docx_dir / f"{base_name}.docx"
            docx_dir.mkdir(parents=True, exist_ok=True)
            file_path = _write_plain_docx(req.content, docx_path, req.options)
            url = f"/AI/word/{file_path.name}"
        elif req.format == "pdf":
            docx_dir.mkdir(parents=True, exist_ok=True)
            pdf_dir.mkdir(parents=True, exist_ok=True)

            docx_path = docx_dir / f"{base_name}.docx"
            _write_plain_docx(req.content, docx_path, req.options)

            converter = OfficialDocumentGenerator()
            pdf_result = converter.save_pdf(docx_path, pdf_dir / f"{base_name}.pdf")
            if not pdf_result:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="PDF导出失败：请安装 LibreOffice 或 unoconv",
                )
            file_path = Path(pdf_result)
            url = f"/AI/pdf/{file_path.name}"
        elif req.format == "txt":
            txt_dir.mkdir(parents=True, exist_ok=True)
            file_path = txt_dir / f"{base_name}.txt"
            file_path.write_text(req.content, encoding="utf-8")
            url = f"/AI/txt/{file_path.name}"
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"暂不支持的导出格式：{req.format}",
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"导出失败：{exc}",
        ) from exc

    expires_at = datetime.now(timezone.utc) + timedelta(days=1)

    return StandardResponse(
        success=True,
        data=DocumentExportData(
            url=url,
            filename=file_path.name,
            size=file_path.stat().st_size,
            expiresAt=expires_at,
        ),
        message="导出成功",
    )



@router.post("/optimize", response_model=StandardResponse[DocumentDataOptimize])
async def document_optimize(
    req: DocumentOptimizeRequest,
    background_tasks: BackgroundTasks,  # ✅ 添加后台任务参数
    svc: DocumentService = Depends(get_document_service),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    POST /document/optimize
    {
        "content": "我们要做好这项工作，效果很好。",
        "optimizationType": "all",
        "customInstruction": "使语气更正式"
    }
    """
    try:
        optimized_text = optimize_document(
            content=req.content,
            optimization_type=req.optimizationType,
            custom_instruction=req.customInstruction
        )
        # print(f"optimized_text: {optimized_text}")
        # 从 LLM 输出中提取 JSON 对象
        json_match = re.search(r'\{[\s\S]*\}', optimized_text)
        if not json_match:
            raise ValueError(f"未能从优化结果中提取JSON对象")
        try:
            document_payload = json.loads(json_match.group())
            lines = []
            for v in document_payload.values():
                if isinstance(v, list):
                    lines.append("\n".join(str(item) for item in v))
                else:
                    lines.append(str(v))

            str_result = "\n".join(lines)
        except json.JSONDecodeError as exc:
            raise ValueError(f"解析生成内容失败：{exc}")


        ai_rate = compute_ai_rate(str_result) if str_result else None
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        docx_dir = Path(os.getenv("DOWNLOAD_DIR", "./generated_documents"))
        pdf_dir = Path(os.getenv("PDF_DIR", "./pdf"))
        docx_path = docx_dir / f"official_document_{timestamp}.docx"
        pdf_path_candidate = pdf_dir / f"official_document_{timestamp}.pdf"

        # 创建生成器
        generator = OfficialDocumentGenerator()
        # 生成文档
        doc = generator.generate(document_payload)
        # 保存为Word
        word_path = generator.save_docx(docx_path)
        # 转换为PDF（可选）
        pdf_path = generator.save_pdf(word_path, pdf_path_candidate)

        word_filename = Path(word_path).name if word_path else None
        pdf_filename = Path(pdf_path).name if pdf_path else None

        docx_preview_path = f"/AI/word/{word_filename}" if word_filename else None
        pdf_preview_path = f"/AI/pdf/{pdf_filename}" if pdf_filename else None
        # print(str_result)

        # ✅ 改为后台任务：不阻塞响应
        optimization_type_map = {
            'grammar': '语法优化',
            'style': '风格优化',
            'clarity': '清晰度优化',
            'logic': '逻辑优化',
            'format': '格式优化',
            'tone': '语气优化',
            'all': '全面优化'
        }
        opt_type_cn = optimization_type_map.get(req.optimizationType, req.optimizationType)
        
        query_record = f"【公文优化】\n优化类型: {opt_type_cn}\n"
        if req.customInstruction:
            query_record += f"自定义指令: {req.customInstruction}\n"
        query_record += f"原文（前200字）: {req.content[:200]}..."
        
        background_tasks.add_task(
            log_conversation_background,
            db=db,
            user_id=current_user["user_id"],
            query=query_record,
            answer=str_result,
            weight=0.8,
            pdf_url=pdf_preview_path,
            word_url=docx_preview_path
        )

        # ✅ 立即返回响应（不等待会话记录完成）
        return StandardResponse(
            success=True,
            data=DocumentDataOptimize(
                content=str_result,
                docxPath=docx_preview_path,
                pdfPath=pdf_preview_path,
                aiRate=ai_rate
            ),
            message="OK"
        )
    except Exception as e:
        logger.error(f"公文优化失败：{e}")
        return StandardResponse(
            success=False,
            data=DocumentDataOptimize(content=""),
            message=f"优化失败：{e}",
        )

@router.post("/ai-rate", response_model=AIRateResponse)
async def document_ai_rate(req: AIRateRequest):
    """Estimate AI-generation likelihood for given text content (0-100).

    Delegates the actual computation to `app.services.ai_rate_service.compute_ai_rate`.
    """
    ai_rate = compute_ai_rate(req.content)
    return AIRateResponse(ai_rate=ai_rate)


class MaterialExtractResponse(BaseModel):
    """素材提取响应。"""
    filename: str
    text: str
    char_count: int


@router.post("/extract-materials", response_model=StandardResponse[list[MaterialExtractResponse]])
async def extract_materials(
    files: list[UploadFile] = File(..., description="素材文件（txt/pdf/docx）"),
    current_user: dict = Depends(get_current_user),
):
    """上传写作素材文件，提取文本内容用于公文生成的参考资料。

    支持格式：.txt / .pdf / .docx
    每个文件返回文件名和提取的文本内容。
    """
    processor = TextProcessor()
    results: list[MaterialExtractResponse] = []
    tmp_path: Optional[Path] = None
    try:
        for f in files:
            suffix = Path(f.filename or "_.txt").suffix.lower()
            if suffix not in (".txt", ".md", ".doc", ".docx", ".pdf"):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"不支持的文件格式: {suffix}，仅支持 .txt / .md / .doc / .docx / .pdf",
                )
            raw = await f.read()
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(raw)
                tmp_path = Path(tmp.name)
            text = processor.extract_text(str(tmp_path))
            results.append(MaterialExtractResponse(
                filename=f.filename or "unknown",
                text=text,
                char_count=len(text),
            ))
            tmp_path.unlink(missing_ok=True)
            tmp_path = None
        return StandardResponse(success=True, data=results, message="素材提取成功")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"素材提取失败: {e}")
        raise HTTPException(status_code=500, detail=f"素材提取失败: {str(e)}")
    finally:
        if tmp_path is not None:
            tmp_path.unlink(missing_ok=True)


# # 使用示例
# if __name__ == "__main__":
#     # 创建生成器
#     generator = OfficialDocumentGenerator()
    
#     # 生成文档
#     doc = generator.generate(EXAMPLE_DATA)
    
#     # 保存为Word
#     generator.save_docx('official_document.docx')
    
#     # 转换为PDF（可选）
#     generator.save_pdf('official_document.docx', 'official_document.pdf')
    
#     print("\n公文生成完成！")
#     print("注意：实际使用时需确保已安装方正字体（FZFangSong-Z02、FZXiaoBiaoSong-B05等）")